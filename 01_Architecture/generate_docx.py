import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import re

def create_document():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    title = doc.add_heading("Project Remora: Architecture & Mathematical Foundations", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_meta = doc.add_paragraph()
    p_meta.add_run("Author: Gerald Corzo (corzogac)\n").bold = True
    p_meta.add_run("Project: Remora — Continuous Wave-Dynamic Latent Inference & Sub-Symbolic Learning\n")
    p_meta.add_run("Date: August 2026\n")
    
    doc.add_heading("1. Executive Overview & Motivation", level=1)
    doc.add_paragraph(
        "Contemporary Large Language Models operate under a discrete symbolic paradigm: every concept, identity, "
        "system rule, and past interaction is translated into tokenized strings. In multi-turn agentic workflows "
        "(such as Janus running at IHE Delft), the system re-tokenizes thousands of tokens of identical persona and "
        "governance definitions before appending a brief user query."
    )
    doc.add_paragraph(
        "Project Remora reformulates inference and adaptation through continuous physical wave mechanics. By treating the "
        "Transformer residual stream as a continuous state-space trajectory z(t, τ), we decouple identity and memory from symbolic text, achieving:\n"
        "1. Zero-Token Latent Ingestion: Pre-computing and hardcoding the persona's standing wave directly into layer biases.\n"
        "2. Sub-Symbolic Lifelong Learning: A lightweight 'pilot-fish' companion network (Remora) that learns directly from the activation wake of the primary model.\n"
        "3. Predictive NVMe Prefetching: Forecasting Mixture-of-Experts (MoE) routing targets to eliminate I/O wait-states in disk-streamed engines like Colibri."
    )
    
    doc.add_heading("2. Mathematical Foundations: The 2D Damped Semantic Wave", level=1)
    doc.add_paragraph(
        "Let τ ∈ [0, L] represent continuous layer depth, and let t ∈ [1, T] represent the discrete sequence token index. "
        "In the continuous depth limit, the residual stream evolution satisfies a non-autonomous Neural Ordinary Differential Equation (Neural ODE):\n\n"
        "   ∂h(t, τ)/∂τ = f_θ(h(t, τ), τ, x_<=t)\n\n"
        "Coupling the depth dynamics τ with the autoregressive context expansion Δt = 1 yields the 2D Damped Semantic Wave Equation:\n\n"
        "   ∂²h(t, τ)/∂τ∂t + γ ∂h(t, τ)/∂τ = D_s ∂²h(t, τ)/∂t² + J_θ(h(t, τ))\n\n"
        "Where γ > 0 is the dissipative damping coefficient (LayerNorm/residual scaling), D_s is the semantic dispersion tensor governing historical back-coupling, and J_θ(h) is the non-linear forcing function (FFN/MoE projections)."
    )
    
    doc.add_heading("3. The Three Wave Regimes & Phase Space Attractors", level=2)
    doc.add_paragraph(
        "1. Ground Potential V_0(h): The neutral language prior where velocity v(t, τ) -> 0.\n"
        "2. Persona Standing Wave V_persona(h): Invariant identity rules establish a stable limit cycle across layer depth. Stored directly as a static latent bias tensor b_persona without tokenization.\n"
        "3. Dialog Traveling Waveform: User queries inject localized kinetic energy, launching a wavefront that propagates through layers and triggers specific MoE expert clusters before dampening back into the limit cycle."
    )
    
    doc.add_heading("4. The Remora Symbiotic Engine (Pilot Fish)", level=1)
    doc.add_paragraph(
        "The Remora companion model rides on the residual stream of the primary heavy Transformer (The Shark). "
        "By continuously sampling activation states and velocity vectors v(t) = h(t) - h(t-1), Remora computes second-order Taylor extrapolations:\n\n"
        "   h_est(t+1, τ) ≈ h(t, τ) + v(t, τ) + 0.5 * a(t, τ)\n\n"
        "This projects directly against MoE router matrices to trigger asynchronous kernel-level prefetch calls (Windows Overlapped I/O / Linux posix_fadvise) before routing layers execute, eliminating NVMe read stalls in Colibri."
    )
    
    doc.add_heading("5. Sub-Symbolic Lifelong Learning", level=1)
    doc.add_paragraph(
        "Instead of maintaining text-based Markdown logs or external RAG vector databases, Remora updates a continuous low-rank dynamic tensor W_remora:\n\n"
        "   ΔW_remora = η * (v(t, τ) ⊗ h(t, τ) - λ W_remora)\n\n"
        "Over time, Remora's internal attractor basins mold to the user's domain vocabulary, reasoning style, and research workflows in pure latent space."
    )

    doc.add_heading("6. Experimental Roadmap", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Experiment"
    hdr_cells[1].text = "Core Focus"
    hdr_cells[2].text = "Key Metric"
    hdr_cells[3].text = "Expected Outcome"
    
    data = [
        ("Exp 1: Latent State Freezing", "Zero-token prompt initialization", "Output Cosine Sim, Tokens Saved", "100% prompt token reduction, <1% output divergence"),
        ("Exp 2: Wave Routing Forecast", "Expert prefetch accuracy", "Top-1 & Top-2 Routing Accuracy", ">75% Top-1 prefetch hit rate; 2x reduction in NVMe read stall"),
        ("Exp 3: Sub-Symbolic Memory", "Continuous activation adaptation", "Trajectory alignment over time", "Organic personalized convergence without text prompts")
    ]
    
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val
            
    doc.save("C:/Users/gco/Dropbox/04-Work/Projects/Remora/01_Architecture/Remora-Architecture.docx")
    print("Successfully generated Remora-Architecture.docx")

if __name__ == "__main__":
    create_document()
